"""模板引擎 — 占位符替换 + 附录表格插入 + 批量生成"""

import copy
import os
import re
from typing import Any, Callable

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 匹配 【xxx】 格式的占位符（中文方括号）
PLACEHOLDER_RE = re.compile(r"【(.+?)】")


def find_placeholders(filepath: str) -> list[str]:
    """扫描模板文件，返回所有【占位符】名称（去重保序）。
    使用逐键探测：对每个候选名称检查 f"【{name}】" 是否实际存在于文本中。
    """
    doc = Document(filepath)
    all_text = "\n".join(_all_texts(doc))

    # 先用正则粗提取所有候选
    candidates = set()
    for m in PLACEHOLDER_RE.finditer(all_text):
        candidates.add(m.group(1))

    # 从候选中筛掉包含【】的（嵌套外壳），只保留叶级占位符
    leaf_candidates = {c for c in candidates if "【" not in c}

    # 验证：只在文本中确实出现 f"【{name}】" 的才算有效
    seen = set()
    result = []
    for name in sorted(leaf_candidates, key=lambda n: all_text.find(f"【{n}】")):
        marker = f"【{name}】"
        if marker in all_text and name not in seen:
            seen.add(name)
            result.append(name)

    # 额外检测：从数据文件列名中可能存在的、在模板中以嵌套方式出现的占位符
    # 如 【协议编号：【协议编号】】 中的 "协议编号"
    for name in candidates:
        inner_marker = f"【{name}】"
        if name not in seen and inner_marker in all_text:
            seen.add(name)
            result.append(name)

    return result


def replace_placeholders(doc: Document, row_data: dict[str, Any]) -> int:
    """替换文档中所有【占位符】，返回替换数量"""
    count = 0

    # 1) 替换段落（正文）
    for para in doc.paragraphs:
        count += _replace_in_paragraph(para, row_data)

    # 2) 替换表格单元格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += _replace_in_paragraph(para, row_data)

    # 3) 替换页眉页脚
    for section in doc.sections:
        for para in section.header.paragraphs:
            count += _replace_in_paragraph(para, row_data)
        for para in section.footer.paragraphs:
            count += _replace_in_paragraph(para, row_data)

    return count


def insert_appendix_table(
    doc: Document,
    placeholder_name: str,
    headers: list[str],
    data_rows: list[list[Any]],
) -> bool:
    """在文档中找到包含【placeholder_name】的段落，删除该段落，在其前一个段落（附录标题）
    后直接插入数据表格。返回是否成功插入。"""
    target_para = None
    for para in doc.paragraphs:
        if f"【{placeholder_name}】" in para.text:
            target_para = para
            break

    if target_para is None:
        return False

    # 找到附录A标题段落（占位符段落的前一个元素）
    body = doc.element.body
    prev_elem = target_para._element.getprevious()
    # 删除占位符段落，使表格紧贴附录A标题
    body.remove(target_para._element)

    # 表格行数 = 1（合并表头）+ 1（列名表头）+ len(data_rows)
    total_rows = len(data_rows) + 2
    anchor = type('obj', (object,), {'_element': prev_elem})()
    table = _insert_table_after_paragraph(doc, anchor, total_rows, len(headers))

    # 找到"内舒拿®"和"妈富隆®"列索引
    neishui_col = None
    mafu_col = None
    for ci, h in enumerate(headers):
        if "内舒拿" in str(h) and neishui_col is None:
            neishui_col = ci
        if "妈富隆" in str(h) and mafu_col is None:
            mafu_col = ci

    # ── 先填充 Row 1（列名表头），必须在 Row 0 设置 gridSpan 之前 ──
    row1_tc_list = list(table.rows[1]._tr.findall(qn("w:tc")))
    for ci, h in enumerate(headers):
        tc = row1_tc_list[ci]
        # 清空并写入
        for p in tc.findall(qn("w:p")):
            tc.remove(p)
        new_p = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_rPr = OxmlElement("w:rPr")
        new_b = OxmlElement("w:b")
        new_rPr.append(new_b)
        new_r.append(new_rPr)
        new_t = OxmlElement("w:t")
        new_t.text = str(h) if h is not None else ""
        new_r.append(new_t)
        new_p.append(new_r)
        tc.append(new_p)

        # 空单元格设置 vMerge continue（和 Row 0 纵向合并）
        is_under_merged = not (neishui_col is not None and mafu_col is not None
                               and neishui_col <= ci <= mafu_col)
        if is_under_merged:
            tcPr = tc.get_or_add_tcPr()
            vm = OxmlElement("w:vMerge")
            tcPr.append(vm)

    # ── 再填充 Row 0（合并表头行）──
    row0_tc_list = list(table.rows[0]._tr.findall(qn("w:tc")))
    for ci in range(len(headers)):
        tc = row0_tc_list[ci]
        # 清空
        for p in tc.findall(qn("w:p")):
            tc.remove(p)

        has_content = (neishui_col is not None and mafu_col is not None
                       and ci == neishui_col)

        if has_content:
            # 有内容：写入标题
            new_p = OxmlElement("w:p")
            new_r = OxmlElement("w:r")
            new_rPr = OxmlElement("w:rPr")
            new_b = OxmlElement("w:b")
            new_rPr.append(new_b)
            new_r.append(new_rPr)
            new_t = OxmlElement("w:t")
            new_t.text = "Q1是否参与分销项目"
            new_r.append(new_t)
            new_p.append(new_r)
            tc.append(new_p)
            # 横向合并
            if mafu_col > neishui_col:
                tcPr = tc.get_or_add_tcPr()
                gs = OxmlElement("w:gridSpan")
                gs.set(qn("w:val"), str(mafu_col - neishui_col + 1))
                tcPr.append(gs)
        else:
            # 空单元格：纵向合并 restart
            tcPr = tc.get_or_add_tcPr()
            vm = OxmlElement("w:vMerge")
            vm.set(qn("w:val"), "restart")
            tcPr.append(vm)

    # ── Row 2+: 数据行 ──
    for ri, row in enumerate(data_rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 2].cells[ci]
            cell.text = str(val) if val is not None else ""

    # ── 删除占位符段落后的空段落（直到遇到附录B等非空段落）──
    _remove_empty_paragraphs_after(doc, anchor)

    # ── 插入分节符：附录A 区域设为横向 ──
    _insert_landscape_section(doc, anchor)

    return True


def generate_single(
    template_path: str,
    row_data: dict[str, Any],
    output_path: str,
    appendix_path: str | None = None,
    appendix_filter_column: str | None = None,
    appendix_match_data_column: str | None = None,
) -> str:
    """生成单份文档（用于预览），返回文件路径。

    参数:
        template_path: 模板 .docx 路径
        row_data: 单行数据 dict
        output_path: 输出文件完整路径（含文件名）
        appendix_path: 附录 Excel 路径 (可选)
        appendix_filter_column: 附录 Excel 中用于筛选的列名
        appendix_match_data_column: 数据文件中用于匹配的列名
    """
    from core.data_loader import load_appendix_data

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document(template_path)

    # ① 先插入附录表格（必须在文本替换之前，否则【附录A】会被清空）
    if appendix_path and appendix_filter_column:
        appendix_headers, appendix_all_rows = load_appendix_data(appendix_path)
        match_value = row_data.get(appendix_match_data_column or "", "")
        if match_value:
            filter_col_idx = None
            for ci, h in enumerate(appendix_headers):
                if h == appendix_filter_column:
                    filter_col_idx = ci
                    break
            if filter_col_idx is not None:
                filtered_rows = [
                    r for r in appendix_all_rows
                    if r[filter_col_idx] is not None
                    and str(r[filter_col_idx]).strip() == str(match_value).strip()
                ]
                insert_appendix_table(doc, "附录A", appendix_headers, filtered_rows)

    # ② 再替换所有文本占位符
    replace_placeholders(doc, row_data)

    doc.save(output_path)
    return output_path


def generate_documents(
    template_path: str,
    data: list[dict[str, Any]],
    output_dir: str,
    appendix_path: str | None = None,
    appendix_filter_column: str | None = None,
    appendix_match_data_column: str | None = None,
    filename_template: str = "{序号}_{经销商名称}",
    progress_callback: Callable[[int, int, str], None] | None = None,
) -> list[str]:
    """批量生成文档。
    返回生成的文件路径列表。

    参数:
        template_path: 模板 .docx 路径
        data: 数据行列表 (list[dict])
        output_dir: 输出目录
        appendix_path: 附录 Excel 路径 (可选)
        appendix_filter_column: 附录 Excel 中用于筛选的列名
        appendix_match_data_column: 数据文件中用于匹配的列名
        filename_template: 文件名模板，支持 {列名} 变量
        progress_callback: 进度回调 (current, total, message)
    """
    os.makedirs(output_dir, exist_ok=True)
    generated_files = []
    total = len(data)

    for i, row in enumerate(data):
        if progress_callback:
            progress_callback(i + 1, total, f"正在生成第 {i + 1}/{total} 份文档...")

        # 生成文件名
        filename = filename_template
        filename = filename.replace("{序号}", str(i + 1).zfill(3))
        for key, val in row.items():
            filename = filename.replace(f"{{{key}}}", str(val))
        # 清理文件名中的非法字符
        filename = re.sub(r'[\\/:*?"<>|]', '_', filename)
        filepath = os.path.join(output_dir, f"{filename}.docx")

        generate_single(
            template_path=template_path,
            row_data=row,
            output_path=filepath,
            appendix_path=appendix_path if appendix_path and os.path.isfile(appendix_path) else None,
            appendix_filter_column=appendix_filter_column or None,
            appendix_match_data_column=appendix_match_data_column or None,
        )
        generated_files.append(filepath)

    return generated_files


# ─────────────────────── 内部函数 ───────────────────────


def _insert_landscape_section(doc: Document, appendix_heading_para):
    """在附录A区域前后插入分节符，使附录A部分为横向。

    预期结构（匹配手动制作的结果文件）：
    - P84 空段落中 sectPr(portrait)  → 签名区与附录A之间
    - P87 空段落中 sectPr(landscape) → 附录A表格与附录B之间
    """
    body = doc.element.body

    # ── 1. 在附录A标题前插入一个空段落，放入 portrait sectPr ──
    heading_elem = appendix_heading_para._element
    new_para = OxmlElement("w:p")
    pPr1 = OxmlElement("w:pPr")
    sectPr1 = OxmlElement("w:sectPr")
    pgSz1 = OxmlElement("w:pgSz")
    pgSz1.set(qn("w:w"), "11907")
    pgSz1.set(qn("w:h"), "16840")
    pgSz1.set(qn("w:code"), "9")
    sectPr1.append(pgSz1)
    pgMar1 = OxmlElement("w:pgMar")
    for attr, val in [("top", "2154"), ("right", "1417"), ("bottom", "993"),
                      ("left", "1701"), ("header", "567"), ("footer", "567"), ("gutter", "0")]:
        pgMar1.set(qn(f"w:{attr}"), val)
    sectPr1.append(pgMar1)
    cols1 = OxmlElement("w:cols")
    cols1.set(qn("w:space"), "720")
    sectPr1.append(cols1)
    pPr1.append(sectPr1)
    new_para.append(pPr1)
    heading_elem.addprevious(new_para)

    # ── 2. 在附录B标题前插入一个空段落，放入 landscape sectPr ──
    found_table = False
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "tbl" and not found_table:
            found_table = True
            continue
        if found_table and tag == "p":
            text = "".join(child.itertext()).strip()
            if text.startswith("附录") and "B" in text:
                new_para2 = OxmlElement("w:p")
                pPr2 = OxmlElement("w:pPr")
                sectPr2 = OxmlElement("w:sectPr")
                pgSz2 = OxmlElement("w:pgSz")
                pgSz2.set(qn("w:w"), "16840")
                pgSz2.set(qn("w:h"), "11907")
                pgSz2.set(qn("w:orient"), "landscape")
                pgSz2.set(qn("w:code"), "9")
                sectPr2.append(pgSz2)
                pgMar2 = OxmlElement("w:pgMar")
                for attr, val in [("top", "1701"), ("right", "2154"), ("bottom", "1417"),
                                  ("left", "993"), ("header", "567"), ("footer", "567"), ("gutter", "0")]:
                    pgMar2.set(qn(f"w:{attr}"), val)
                sectPr2.append(pgMar2)
                cols2 = OxmlElement("w:cols")
                cols2.set(qn("w:space"), "720")
                sectPr2.append(cols2)
                docGrid = OxmlElement("w:docGrid")
                docGrid.set(qn("w:linePitch"), "272")
                sectPr2.append(docGrid)
                pPr2.append(sectPr2)
                new_para2.append(pPr2)
                child.addprevious(new_para2)
                break


def _display_width(s: str) -> int:
    """计算字符串显示宽度：中日韩/全角字符算 2，ASCII 算 1。"""
    width = 0
    for ch in s:
        cp = ord(ch)
        if (
            (0x4E00 <= cp <= 0x9FFF)
            or (0x3000 <= cp <= 0x303F)
            or (0xFF00 <= cp <= 0xFFEF)
            or (0x3400 <= cp <= 0x4DBF)
            or (0x2E80 <= cp <= 0x2EFF)
            or (0xF900 <= cp <= 0xFAFF)
        ):
            width += 2
        else:
            width += 1
    return width


def _fix_party_alignment(doc: Document):
    """修正乙方行中（简称：的对齐，使其与甲方行对齐。"""
    jia_col = None
    yi_para = None

    for para in doc.paragraphs:
        text = para.text
        if "（简称：" in text:
            if "甲方" in text and jia_col is None:
                jia_col = _display_width(text[:text.index("（简称：")])
            elif "乙方" in text:
                yi_para = para

    if jia_col is None or yi_para is None:
        return

    text = yi_para.text
    marker = "（简称："
    if marker not in text:
        return

    idx = text.index(marker)
    current_col = _display_width(text[:idx])
    if current_col == jia_col:
        return

    diff = jia_col - current_col  # 正值=需要加空格，负值=需要减空格
    prefix = text[:idx]
    suffix = text[idx:]

    if diff > 0:
        # 需要加宽：追加空格
        prefix = prefix + " " * diff
    elif diff < 0:
        # 需要缩窄：从末尾删除多余的 ASCII 空格
        remove_count = -diff  # 需要移除的显示列数（每个空格=1列）
        prefix = prefix.rstrip(" ")
        # 加回正确数量的空格
        remaining_spaces = len(text[:idx]) - len(prefix) - remove_count
        if remaining_spaces > 0:
            prefix = prefix + " " * remaining_spaces

    new_text = prefix + suffix

    # 写回段落
    if yi_para.runs:
        for run in yi_para.runs:
            run.text = ""
        yi_para.runs[0].text = new_text
    else:
        run = yi_para.add_run(new_text)


def _remove_empty_paragraphs_after(doc: Document, anchor_para):
    """删除 anchor_para 之后、到下一个非空段落之间的所有空段落。"""
    body = doc.element.body
    anchor_elem = anchor_para._element

    # 找到 anchor 的位置
    found = False
    to_remove = []
    for child in body:
        if child is anchor_elem:
            found = True
            continue
        if not found:
            continue

        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "tbl":
            # 跳过刚插入的表格
            continue
        if tag == "p":
            text = "".join(child.itertext()).strip()
            if text == "":
                to_remove.append(child)
            else:
                # 遇到非空段落，停止
                break
        elif tag == "sectPr":
            break

    for elem in to_remove:
        body.remove(elem)


def _merge_cells_horizontal(row, start_col: int, end_col: int):
    """合并表格行中 start_col 到 end_col 的单元格（水平合并）。"""
    for ci in range(start_col, end_col + 1):
        tc = row.cells[ci]._tc
        tcPr = tc.get_or_add_tcPr()
        if ci == start_col:
            gridSpan = OxmlElement("w:gridSpan")
            gridSpan.set(qn("w:val"), str(end_col - start_col + 1))
            tcPr.append(gridSpan)
        else:
            # 被合并的单元格标记为 continue
            tcPr.append(OxmlElement("w:hideMark"))


def _all_texts(doc: Document) -> list[str]:
    """收集文档中所有文本（正文 + 表格 + 页眉页脚）"""
    texts = []
    for para in doc.paragraphs:
        texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    texts.append(para.text)
    for section in doc.sections:
        for para in section.header.paragraphs:
            texts.append(para.text)
        for para in section.footer.paragraphs:
            texts.append(para.text)
    return texts


def _replace_in_paragraph(para, row_data: dict[str, Any]) -> int:
    """在段落中替换占位符。使用逐键字符串替换，正确处理嵌套占位符。"""
    full_text = para.text
    if "【" not in full_text:
        return 0

    new_text = full_text
    count = 0
    for key, val in row_data.items():
        placeholder = f"【{key}】"
        if placeholder in new_text:
            occurrences = new_text.count(placeholder)
            new_text = new_text.replace(placeholder, str(val))
            count += occurrences

    if count == 0:
        return 0

    # 保留第一个 run 的格式属性
    if para.runs:
        first_run = para.runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        bold = first_run.font.bold
        italic = first_run.font.italic
    else:
        font_name = font_size = bold = italic = None

    # 清空原有 runs，重新写入
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        run = para.add_run(new_text)
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic

    return count


def _insert_table_after_paragraph(doc: Document, para, rows: int, cols: int):
    """在指定段落后插入一个表格"""
    # 创建表格
    table = doc.add_table(rows=rows, cols=cols, style="Table Grid")

    # 将表格 XML 移到目标段落后面
    table_element = table._tbl
    para_element = para._element
    para_element.addnext(table_element)

    # 设置表格自动适配宽度
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    return table
