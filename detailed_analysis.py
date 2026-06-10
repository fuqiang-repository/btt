#!/usr/bin/env python3
import re
from docx import Document
from docx.table import Table

def analyze_document_structure(docx_file):
    """Analyze the complete document structure"""
    print(f"\nAnalyzing document structure for: {docx_file}")
    print("=" * 80)
    
    doc = Document(docx_file)
    
    # Document properties
    print(f"Document Properties:")
    print(f"  Core Properties:")
    for prop in doc.core_properties:
        if prop.value:
            print(f"    {prop}: {prop.value}")
    
    # Sections analysis
    print(f"\nSections: {len(doc.sections)}")
    for i, section in enumerate(doc.sections):
        print(f"  Section {i+1}:")
        print(f"    Page width: {section.page_width}")
        print(f"    Page height: {section.page_height}")
        print(f    Top margin: {section.top_margin}")
        print(f    Bottom margin: {section.bottom_margin}")
        print(f    Left margin: {section.left_margin}")
        print(f    Right margin: {section.right_margin}")
        
        # Headers and footers
        header = section.header
        footer = section.footer
        if header.paragraphs and any(p.text.strip() for p in header.paragraphs):
            print(f"    Header text: {header.paragraphs[0].text.strip()}")
        if footer.paragraphs and any(p.text.strip() for p in footer.paragraphs):
            print(f"    Footer text: {footer.paragraphs[0].text.strip()}")
    
    # Paragraphs with styles
    print(f"\nParagraphs with styles:")
    placeholders = []
    styled_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            style_name = para.style.name if para.style else "Normal"
            text = para.text.strip()
            
            # Check for placeholders
            if '【' in text and '】' in text:
                matches = re.findall(r'【([^】]+)】', text)
                placeholders.extend(matches)
                print(f"  [P{i+1:3d}] Style: {style_name:20} | Placeholder: {matches[0]}")
            
            # Show styled paragraphs
            if style_name not in ['Normal', 'Default Paragraph Font'] and not style_name.startswith('Heading'):
                styled_paragraphs.append((i+1, style_name, text[:50]))
    
    if styled_paragraphs:
        print(f"\nStyled paragraphs (excluding Normal and Headings):")
        for idx, style, preview in styled_paragraphs:
            print(f"  [{idx:3d}] {style:20} | {preview}...")
    
    # Tables analysis
    print(f"\nTables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1}:")
        print(f"  Dimensions: {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Check for placeholders in table cells
        table_placeholders = []
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text and '【' in cell_text and '】' in cell_text:
                    matches = re.findall(r'【([^】]+)】', cell_text)
                    table_placeholders.extend(matches)
                    print(f"    Cell [{row_idx+1},{col_idx+1}]: {matches[0]}")
        
        if table_placeholders:
            print(f"  Total placeholders in table: {len(table_placeholders)}")
    
    return {
        'placeholders': list(set(placeholders)),
        'total_paragraphs': len([p for p in doc.paragraphs if p.text.strip()]),
        'total_tables': len(doc.tables),
        'sections': len(doc.sections)
    }

def find_all_occurrences(docx_file, search_text):
    """Find all occurrences of specific text in document"""
    print(f"\nSearching for occurrences of: {search_text}")
    print("-" * 50)
    
    doc = Document(docx_file)
    occurrences = []
    
    # Search in paragraphs
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            occurrences.append(('paragraph', i+1, para.text.strip()))
    
    # Search in tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if search_text in cell.text:
                    occurrences.append(('table', table_idx+1, row_idx+1, col_idx+1, cell.text.strip()))
    
    if occurrences:
        for occ in occurrences:
            if occ[0] == 'paragraph':
                print(f"  Paragraph {occ[1]}: {occ[2][:80]}...")
            else:
                print(f"  Table {occ[1]}[{occ[2]},{occ[3]}]: {occ[4][:80]}...")
    else:
        print("  No occurrences found")

def compare_placeholder_content(template_file, result_file):
    """Compare the content of placeholders between template and result"""
    print(f"\nComparing placeholder content")
    print("=" * 80)
    
    template_doc = Document(template_file)
    result_doc = Document(result_file)
    
    template_placeholders = {}
    result_placeholders = {}
    
    # Extract placeholders from template
    for para in template_doc.paragraphs:
        text = para.text
        if '【' in text and '】' in text:
            matches = re.findall(r'【([^】]+)】', text)
            for ph in matches:
                template_placeholders[ph] = text
    
    # Extract placeholders from result
    for para in result_doc.paragraphs:
        text = para.text
        if '【' in text and '】' in text:
            matches = re.findall(r'【([^】]+)】', text)
            for ph in matches:
                result_placeholders[ph] = text
    
    # Compare
    print("Placeholders in template:")
    for ph in sorted(template_placeholders.keys()):
        print(f"  {ph}: {template_placeholders[ph][:80]}...")
    
    print("\nPlaceholders in result:")
    for ph in sorted(result_placeholders.keys()):
        print(f"  {ph}: {result_placeholders[ph][:80]}...")
    
    # Find differences
    unique_to_template = set(template_placeholders.keys()) - set(result_placeholders.keys())
    unique_to_result = set(result_placeholders.keys()) - set(template_placeholders.keys())
    
    if unique_to_template:
        print(f"\nPlaceholders only in template:")
        for ph in unique_to_template:
            print(f"  - {ph}")
    
    if unique_to_result:
        print(f"\nPlaceholders only in result:")
        for ph in unique_to_result:
            print(f"  - {ph}")

def main():
    template_file = "/Users/fuqiang/PycharmProjects/TemplateForge/分销协议制作/D-欧加隆2026年零售重点客户年度销售合作协议补充协议-季度机会品分销项目_ 模板.docx"
    result_file = "/Users/fuqiang/PycharmProjects/TemplateForge/分销协议制作/D-欧加隆2026年零售重点客户年度销售合作协议补充协议-季度机会品分销项目_.docx"
    
    # Analyze document structure
    template_analysis = analyze_document_structure(template_file)
    result_analysis = analyze_document_structure(result_file)
    
    # Compare placeholder content
    compare_placeholder_content(template_file, result_file)
    
    # Search for specific placeholders
    print(f"\nDetailed placeholder search")
    print("=" * 80)
    
    # Common placeholders to search for
    common_placeholders = ['经销商名称', '附录A', '欧加隆盖章处', '经销商盖章处']
    
    for ph in common_placeholders:
        find_all_occurrences(template_file, ph)
        find_all_occurrences(result_file, ph)
    
    # Summary
    print(f"\nSUMMARY")
    print("=" * 80)
    print(f"Template document:")
    print(f"  - Total paragraphs: {template_analysis['total_paragraphs']}")
    print(f"  - Tables: {template_analysis['tables']}")
    print(f"  - Sections: {template_analysis['sections']}")
    print(f"  - Unique placeholders: {len(template_analysis['placeholders'])}")
    
    print(f"\nResult document:")
    print(f"  - Total paragraphs: {result_analysis['total_paragraphs']}")
    print(f"  - Tables: {result_analysis['tables']}")
    print(f"  - Sections: {result_analysis['sections']}")
    print(f"  - Unique placeholders: {len(result_analysis['placeholders']}")

if __name__ == "__main__":
    main()
