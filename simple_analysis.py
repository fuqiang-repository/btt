#!/usr/bin/env python3
import re
from docx import Document

def analyze_docx(docx_file, name):
    print(f"\n{'='*60}")
    print(f"Analyzing: {name}")
    print(f"{'='*60}")
    
    doc = Document(docx_file)
    
    # Basic stats
    total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
    total_tables = len(doc.tables)
    total_sections = len(doc.sections)
    
    print(f"\nDocument Statistics:")
    print(f"  - Total paragraphs: {total_paragraphs}")
    print(f"  - Total tables: {total_tables}")
    print(f"  - Total sections: {total_sections}")
    
    # Find all 【】 placeholders
    all_placeholders = []
    paragraph_occurrences = []
    table_occurrences = []
    
    # Check paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and '【' in text:
            matches = re.findall(r'【([^】]+)】', text)
            for match in matches:
                all_placeholders.append(match)
                paragraph_occurrences.append((i+1, match, text[:100]))
    
    # Check tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text and '【' in cell_text:
                    matches = re.findall(r'【([^】]+)】', cell_text)
                    for match in matches:
                        all_placeholders.append(match)
                        table_occurrences.append((table_idx+1, row_idx+1, col_idx+1, match, cell_text[:100]))
    
    # Show unique placeholders
    unique_placeholders = sorted(set(all_placeholders))
    print(f"\nUnique Placeholders ({len(unique_placeholders)}):")
    for i, ph in enumerate(unique_placeholders, 1):
        print(f"  {i:2d}. {ph}")
    
    # Show paragraph occurrences
    if paragraph_occurrences:
        print(f"\nPlaceholders in paragraphs:")
        for idx, match, context in paragraph_occurrences:
            print(f"  P{idx}: 【{match}】 - {context}...")
    
    # Show table occurrences
    if table_occurrences:
        print(f"\nPlaceholders in tables:")
        for table_idx, row_idx, col_idx, match, context in table_occurrences:
            print(f"  T{table_idx}[{row_idx},{col_idx}]: 【{match}】 - {context}...")
    
    return {
        'total_paragraphs': total_paragraphs,
        'total_tables': total_tables,
        'total_sections': total_sections,
        'unique_placeholders': unique_placeholders,
        'total_placeholders': len(all_placeholders),
        'paragraph_occurrences': paragraph_occurrences,
        'table_occurrences': table_occurrences
    }

def compare_files(template_file, result_file):
    print(f"\n{'='*60}")
    print("COMPARING TEMPLATE AND RESULT")
    print(f"{'='*60}")
    
    # Get all text from both files
    def get_all_text(doc):
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        return text
    
    template_doc = Document(template_file)
    result_doc = Document(result_file)
    
    template_text = get_all_text(template_doc)
    result_text = get_all_text(result_doc)
    
    # Find text differences
    template_only = set(template_text) - set(result_text)
    result_only = set(result_text) - set(template_text)
    
    print(f"\nText only in template ({len(template_only)} items):")
    for i, text in enumerate(sorted(template_only), 1):
        print(f"  {i}. {text[:80]}...")
    
    print(f"\nText only in result ({len(result_only)} items):")
    for i, text in enumerate(sorted(result_only), 1):
        print(f"  {i}. {text[:80]}...")

def main():
    template_file = "/Users/fuqiang/PycharmProjects/TemplateForge/分销协议制作/D-欧加隆2026年零售重点客户年度销售合作协议补充协议-季度机会品分销项目_ 模板.docx"
    result_file = "/Users/fuqiang/PycharmProjects/TemplateForge/分销协议制作/D-欧加隆2026年零售重点客户年度销售合作协议补充协议-季度机会品分销项目_.docx"
    
    # Analyze both files
    template_data = analyze_docx(template_file, "Template")
    result_data = analyze_docx(result_file, "Result")
    
    # Compare files
    compare_files(template_file, result_file)
    
    # Summary
    print(f"\n{'='*60}")
    print("ANALYSIS SUMMARY")
    print(f"{'='*60}")
    
    print(f"\nTemplate File:")
    print(f"  - Paragraphs: {template_data['total_paragraphs']}")
    print(f"  - Tables: {template_data['total_tables']}")
    print(f"  - Sections: {template_data['total_sections']}")
    print(f"  - Unique placeholders: {len(template_data['unique_placeholders'])}")
    print(f"  - Total placeholder occurrences: {template_data['total_placeholders']}")
    
    print(f"\nResult File:")
    print(f"  - Paragraphs: {result_data['total_paragraphs']}")
    print(f"  - Tables: {result_data['total_tables']}")
    print(f"  - Sections: {result_data['total_sections']}")
    print(f"  - Unique placeholders: {len(result_data['unique_placeholders'])}")
    print(f"  - Total placeholder occurrences: {result_data['total_placeholders']}")
    
    # Find differences in placeholders
    template_only_ph = set(template_data['unique_placeholders']) - set(result_data['unique_placeholders'])
    result_only_ph = set(result_data['unique_placeholders']) - set(template_data['unique_placeholders'])
    
    if template_only_ph:
        print(f"\nPlaceholders only in template:")
        for ph in sorted(template_only_ph):
            print(f"  - {ph}")
    
    if result_only_ph:
        print(f"\nPlaceholders only in result:")
        for ph in sorted(result_only_ph):
            print(f"  - {ph}")

if __name__ == "__main__":
    main()
