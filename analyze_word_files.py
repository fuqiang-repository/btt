#!/usr/bin/env python3
import re
import sys
import os
from docx import Document
from docx.shared import Inches

def extract_text_and_placeholders(docx_file):
    """Extract all text content and identify placeholders from a Word document"""
    print(f"\n{'='*60}")
    print(f"Analyzing file: {os.path.basename(docx_file)}")
    print(f"{'='*60}")
    
    try:
        doc = Document(docx_file)
        
        # Extract all text
        all_text = []
        paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                all_text.append(text)
                paragraphs.append((i+1, text))
        
        # Extract placeholders with 【】 style
        placeholders = []
        for para_text in all_text:
            # Find 【】 style placeholders
            matches = re.findall(r'【([^】]+)】', para_text)
            if matches:
                placeholders.extend(matches)
        
        # Find other common placeholder patterns
        other_placeholders = []
        for para_text in all_text:
            # Find ${var} style placeholders
            matches = re.findall(r'\$\{([^}]+)\}', para_text)
            other_placeholders.extend(matches)
            
            # Find {{var}} style placeholders
            matches = re.findall(r'\{\{([^}]+)\}\}', para_text)
            other_placeholders.extend(matches)
            
            # Find [var] style placeholders
            matches = re.findall(r'\[([^\]]+)\]', para_text)
            other_placeholders.extend(matches)
        
        # Get document structure info
        num_sections = len(doc.sections)
        tables = []
        for table in doc.tables:
            table_data = []
            for i, row in enumerate(table.rows):
                row_data = []
                for j, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            tables.append((len(tables) + 1, table_data))
        
        # Print results
        print(f"\nDocument Statistics:")
        print(f"  - Total paragraphs: {len(paragraphs)}")
        print(f"  - Number of sections: {num_sections}")
        print(f"  - Number of tables: {len(tables)}")
        
        print(f"\n{'【】 Style Placeholders:':<30} {len(placeholders)} found")
        if placeholders:
            for i, ph in enumerate(sorted(set(placeholders)), 1):
                print(f"  {i:2d}. {ph}")
        
        if other_placeholders:
            print(f"\n{'Other Placeholders:':<30} {len(other_placeholders)} found")
            for i, ph in enumerate(sorted(set(other_placeholders)), 1):
                print(f"  {i:2d}. {ph}")
        else:
            print(f"\n{'Other Placeholders:':<30} None found")
        
        # Show first few paragraphs for context
        print(f"\nFirst 10 paragraphs:")
        for i, (idx, text) in enumerate(paragraphs[:10], 1):
            print(f"  {idx:3d}: {text[:100]}{'...' if len(text) > 100 else ''}")
        
        # Show table information
        if tables:
            print(f"\nTables:")
            for table_idx, table_data in tables:
                print(f"  Table {table_idx}: {len(table_data)} rows")
                if len(table_data) > 0:
                    print(f"    Columns: {len(table_data[0])}")
                    print(f"    First row: {table_data[0]}")
        
        return {
            'file': os.path.basename(docx_file),
            'paragraphs': len(paragraphs),
            'sections': num_sections,
            'tables': len(tables),
            'placeholders': sorted(set(placeholders)),
            'other_placeholders': sorted(set(other_placeholders)),
            'first_paragraphs': paragraphs[:10],
            'tables_info': [(idx, len(data), len(data[0]) if data else 0) for idx, data in tables]
        }
    
    except Exception as e:
        print(f"Error reading {docx_file}: {str(e)}")
        return None

def compare_documents(template_file, result_file):
    """Compare two Word documents to see differences"""
    print(f"\n{'='*60}")
    print("Comparing Documents")
    print(f"{'='*60}")
    
    template_doc = Document(template_file)
    result_doc = Document(result_file)
    
    template_text = [p.text.strip() for p in template_doc.paragraphs if p.text.strip()]
    result_text = [p.text.strip() for p in result_doc.paragraphs if p.text.strip()]
    
    print(f"Template: {len(template_text)} paragraphs")
    print(f"Result: {len(result_text)} paragraphs")
    
    # Find differences
    template_only = set(template_text) - set(result_text)
    result_only = set(result_text) - set(template_text)
    
    if template_only:
        print(f"\nText only in template ({len(template_only)} items):")
        for i, text in enumerate(sorted(template_only), 1):
            print(f"  {i}. {text[:80]}{'...' if len(text) > 80 else ''}")
    
    if result_only:
        print(f"\nText only in result ({len(result_only)} items):")
        for i, text in enumerate(sorted(result_only), 1):
            print(f"  {i}. {text[:80]}{'...' if len(text) > 80 else ''}")

def main():
    # File paths
    template_file = "/Users/fuqiang/PycharmProjects/TemplateForge/分销协议制作/D-欧加隆2026年零售重点客户年度销售合作协议补充协议-季度机会品分销项目_ 模板.docx"
    result_file = "/Users/fuqiang/PycharmProjects/TemplateForge/分销协议制作/D-欧加隆2026年零售重点客户年度销售合作协议补充协议-季度机会品分销项目_.docx"
    
    # Analyze both files
    template_data = extract_text_and_placeholders(template_file)
    result_data = extract_text_and_placeholders(result_file)
    
    # Compare documents
    compare_documents(template_file, result_file)
    
    # Summary
    print(f"\n{'='*60}")
    print("ANALYSIS SUMMARY")
    print(f"{'='*60}")
    
    if template_data and result_data:
        print(f"\nTemplate File: {template_data['file']}")
        print(f"  - Total placeholders: {len(template_data['placeholders'])}")
        print(f"  - Other placeholders: {len(template_data['other_placeholders'])}")
        
        print(f"\nResult File: {result_data['file']}")
        print(f"  - Total placeholders: {len(result_data['placeholders'])}")
        print(f"  - Other placeholders: {len(result_data['other_placeholders'])}")
        
        # Find filled placeholders
        if template_data['placeholders']:
            filled_placeholders = set(template_data['placeholders']) - set(result_data['placeholders'])
            print(f"\nFilled Placeholders ({len(filled_placeholders)}):")
            for ph in sorted(filled_placeholders):
                print(f"  - {ph}")

if __name__ == "__main__":
    main()
